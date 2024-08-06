from transformers import BertTokenizer, BertModel
import torch
import torch.nn.functional as F

import os
import docx
from docx import Document


tokenizer = BertTokenizer.from_pretrained(
    'shibing624/text2vec-bge-large-chinese')  # shibing624/text2vec-bge-large-chinese     shibing624/text2vec-base-chinese   sentence-transformers/all-MiniLM-L6-v2   # model_max_length=xxx
model = BertModel.from_pretrained(
    'shibing624/text2vec-bge-large-chinese')  # 效果                  >            效果                   >


# Mean Pooling - Take attention mask into account for correct averaging
def mean_pooling(model_output, attention_mask):
    token_embeddings = model_output[0]  # First element of model_output contains all token embeddings
    input_mask_expanded = attention_mask.unsqueeze(-1).expand(token_embeddings.size()).float()
    return torch.sum(token_embeddings * input_mask_expanded, 1) / torch.clamp(input_mask_expanded.sum(1), min=1e-9)


def heading2vec(heading_sentences):
    # Load model from HuggingFace Hub

    sentences = heading_sentences

    # Tokenize sentences
    encoded_input = tokenizer(sentences, padding=True, truncation=True, return_tensors='pt')

    # Compute token embeddings
    with torch.no_grad():
        model_output = model(**encoded_input)
    # Perform pooling. In this case, mean pooling.
    sentence_embeddings = mean_pooling(model_output, encoded_input['attention_mask'])
    return sentence_embeddings


# 比较向量相似度
def compute_similarity(input_text, sentence_embedding):
    # Tokenize input text
    encoded_input = tokenizer([input_text], padding=True, truncation=True, return_tensors='pt')

    # Compute token embeddings
    with torch.no_grad():
        model_output = model(**encoded_input)
    # Perform pooling. In this case, mean pooling.
    input_embedding = mean_pooling(model_output, encoded_input['attention_mask'])

    # Compute cosine similarity
    similarities = F.cosine_similarity(input_embedding, sentence_embedding)
    return similarities


# 判断哪个章节标题与目标内容更相似
def check_similarity(similarities, confidence):
    index = []
    for i in range(len(similarities)):
        if float(similarities[i]) > confidence:
            index.append(i)
    return index


def list_headings(doc_path):
    # 打开docx文档
    doc = Document(doc_path)

    # 用于存储所有标题
    headings = []

    for paragraph in doc.paragraphs:
        # 检查段落是否是标题
        if paragraph.style.name.startswith('Heading'):
            if paragraph.text != '':
                headings.append(paragraph.text)

    return headings


def print_heading(index, heading):
    match_heading = []
    for it in index:
        match_heading.append(heading[it])
    return match_heading


# 查询某标题的内容在文档中的段落位置
def find_heading_range(docx_path, heading_text):
    # 打开文档
    doc = Document(docx_path)

    start_index = None
    end_index = None
    in_heading = False

    # 遍历所有段落
    for index, paragraph in enumerate(doc.paragraphs):
        # 检查段落的样式
        if paragraph.style.name.startswith('Heading'):
            if in_heading:
                # 如果之前已经在一个标题区域中，更新结束位置
                end_index = index - 1
                break

            # 如果当前段落是目标标题
            if heading_text in paragraph.text:
                start_index = index
                in_heading = True

    # 如果最后一个标题区域没有结束位置，则设为文档末尾
    if in_heading and end_index is None:
        end_index = len(doc.paragraphs) - 1

    return (int(start_index), int(end_index))


# 遍历表格在文档中的段落位置
def get_table_paragraph_positions(docx_path):
    document = Document(docx_path)
    table_positions = []
    para_counter = 0

    for block in document.element.body:
        if block.tag.endswith('tbl'):
            table_positions.append(para_counter)
        elif block.tag.endswith('p'):
            para_counter += 1

    return table_positions


# 列表逐行保存
def table_2_text(number, docx_path):
    doc = docx.Document(docx_path)
    tables = doc.tables
    tb = tables[number]
    table2text = []

    tb_rows = tb.rows
    for i in range(len(tb_rows)):
        row_data = []
        row_cells = tb_rows[i].cells
        # 读取每一行单元格内容
        for cell in row_cells:
            # 单元格内容
            row_data.append(cell.text)
        table2text.append(row_data)

    return table2text


# 计算表格相对与标题下内容的相对位置
def check_paragraphs(head_para, table_para):
    insert_para = table_para - head_para

    return insert_para


# 根据heading_text 拆分文档并另存新的文档
def extract_content_by_heading(doc_path, heading_text, output_path):
    doc = Document(doc_path)

    # 用于存储要保存的段落
    new_doc = Document()
    save_content = False

    for paragraph in doc.paragraphs:
        # 检查段落是否是标题
        if paragraph.style.name.startswith('Heading'):
            # 如果找到指定标题，开始保存后续内容
            if paragraph.text == heading_text:
                save_content = True
            else:
                save_content = False

        # 如果当前段落需要保存，则将其添加到新文档
        if save_content and not paragraph.style.name.startswith('Heading'):
            new_doc.add_paragraph(paragraph.text, style=paragraph.style)

    # 保存新文档
    new_doc.save(output_path)


# 将文本化的表格 插入拆分后的文档
def insert_text_in_paragraph(output_path, paragraph_index, table2text):
    doc = Document(output_path)
    text_string = ''
    # 检查是否有足够的段落
    if paragraph_index >= len(doc.paragraphs):
        new_paragraph = doc.add_paragraph("")
        # 将新段落插入到文档的开始
        doc.paragraphs.insert(0, new_paragraph)

    for i in range(len(table2text)):
        text_string = text_string + str(table2text[i])
    # 获取指定的段落
    paragraph = doc.paragraphs[paragraph_index]

    # 插入文本到指定段落
    paragraph.add_run(text_string)

    # 获取文件名和路径
    dir_name, base_name = os.path.split(output_path)

    # new_file_name = f'updated_{base_name}'   #保留切分文件并另存新的文件
    new_file_name = f'{base_name}'  # 在切分文件基础上保存
    new_file_path = os.path.join(dir_name, new_file_name)

    # 保存文档
    doc.save(new_file_path)
    print(f"文本已插入到第 {paragraph_index} 段落，并保存为 '{new_file_path}'。")


# 计算表格属于哪个章节标题,并插入表格
def insert_table(list, min, max, heading, bid_path, output_path):
    for i in range(len(list)):
        if list[i] >= int(min):
            if list[i] <= int(max):
                table_2_text(i, bid_path)
                print(table_2_text(i, bid_path))
                extract_content_by_heading(bid_path, heading, output_path)
                s = check_paragraphs(int(min), list[i])
                insert_text_in_paragraph(output_path, s, table_2_text(i, bid_path))


#预处理文档，拿来即用
def preprocess_docx(bid_path, out_path, confidence, Match_text, ):
    headings = list_headings(bid_path)
    print(headings)
    compute_s = compute_similarity(Match_text, heading2vec(headings))
    print(compute_s)
    check_s = check_similarity(compute_s, confidence)
    heading = print_heading(check_s, headings)

    for i in range(len(heading)):
        min, max = find_heading_range(bid_path, heading[i])
        list = get_table_paragraph_positions(bid_path)
        extract_content_by_heading(bid_path, heading[i], out_path + heading[i] + '.docx')
        insert_table(list, min, max, heading[i], bid_path, out_path + heading[i] + '.docx')
